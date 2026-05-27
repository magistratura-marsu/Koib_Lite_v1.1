# -*- coding: utf-8 -*-
"""
Koib-V-4.5 — Модуль парсинга документов
==========================================
Извлечение структурированных элементов из PDF (PyMuPDF) и DOCX.
OCR только через Tesseract (лёгкий, без загрузки модели в Python).

Классы:
  - DocumentElement — структурированный элемент документа

Функции:
  - parse_pdf   — парсинг PDF-файла
  - parse_docx  — парсинг DOCX-файла
"""
import io
import re
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field, asdict

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

from .utils import (
    clean_text, text_hash, detect_model_in_text,
    detect_model_from_filename, find_figure_caption,
    extract_headings, estimate_tokens, generate_unique_id,
)
from config import (
    OCR_DPI, OCR_MIN_TEXT_CHARS, MIN_IMAGE_WIDTH,
    MIN_IMAGE_HEIGHT, PARSING_ENGINE, FIGURES_DIR,
)

logger = logging.getLogger("koib.parsing")


# ═══════════════════════════════════════════════════════════════
# Структура элемента документа
# ═══════════════════════════════════════════════════════════════
@dataclass
class DocumentElement:
    """
    Структурированный элемент документа.

    Attributes:
        content:     Текстовое содержимое элемента
        element_type: Тип элемента (text, table, formula, figure, heading)
        source:      Имя файла-источника
        page:        Номер страницы (для PDF)
        heading:     Заголовок раздела, к которому принадлежит элемент
        model:       Название модели/устройства (если обнаружено)
        element_id:  Уникальный идентификатор элемента
        metadata:    Дополнительные метаданные
    """
    content: str
    element_type: str  # text, table, formula, figure, heading
    source: str = ""
    page: int = 0
    heading: str = ""
    model: str = "unknown"
    element_id: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if not self.element_id:
            self.element_id = text_hash(
                f"{self.source}:{self.page}:{self.element_type}:{self.content[:200]}"
            )

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    @property
    def is_structured(self) -> bool:
        """Является ли элемент структурированным (таблица/формула/рисунок)."""
        return self.element_type in ("table", "formula", "figure")


# ═══════════════════════════════════════════════════════════════
# Вспомогательные функции
# ═══════════════════════════════════════════════════════════════
def _expand_rect(rect: fitz.Rect, margin: float) -> fitz.Rect:
    """Расширить прямоугольник на заданный отступ."""
    return fitz.Rect(
        rect.x0 - margin,
        rect.y0 - margin,
        rect.x1 + margin,
        rect.y1 + margin,
    )


def _is_scanned_page(page: fitz.Page, min_chars: int = OCR_MIN_TEXT_CHARS) -> bool:
    """
    Определить, является ли страница сканом.
    Если текста мало и есть крупное изображение — страница сканированная.
    """
    text = page.get_text("text").strip()
    if len(text) >= min_chars:
        return False

    images = page.get_images(full=True)
    if not images:
        return len(text) < min_chars

    page_area = page.rect.width * page.rect.height
    for img_info in images:
        try:
            xref = img_info[0]
            base_image = page.parent.extract_image(xref)
            if not base_image:
                continue
            img = Image.open(io.BytesIO(base_image["image"]))
            if img.width * img.height / page_area > 0.8:
                return True
        except Exception:
            continue
    return True


def _ocr_image(image_pil: Image.Image, lang: str = "rus+eng") -> str:
    """
    Распознать текст на изображении через Tesseract OCR.
    Tesseract работает как внешний процесс (нулевое потребление RAM
    в Python). В отличие от EasyOCR, не загружает нейросеть в память.
    """
    if image_pil is None:
        return ""
    try:
        import pytesseract
        text = clean_text(
            pytesseract.image_to_string(image_pil, lang=lang, config="--psm 6")
        )
        if len(text) >= 30:
            return text
    except Exception as exc:
        logger.debug(f"Tesseract OCR error: {exc}")
    return ""


def _extract_tables_from_page(page: fitz.Page) -> List[Dict[str, Any]]:
    """
    Извлечь таблицы со страницы PDF через PyMuPDF.
    Возвращает список словарей с ключами: text, rows, cols, bbox.
    """
    tables = []
    try:
        tab_finder = page.find_tables()
        for tab in tab_finder:
            try:
                df = tab.to_pandas()
                markdown = df.to_markdown(index=False)
                rows, cols = df.shape
                tables.append({
                    "text": markdown,
                    "num_rows": rows,
                    "num_cols": cols,
                    "bbox": tuple(tab.bbox) if hasattr(tab, "bbox") else (0, 0, 0, 0),
                })
            except Exception as exc:
                logger.debug(f"Ошибка конвертации таблицы: {exc}")
    except Exception as exc:
        logger.debug(f"Ошибка поиска таблиц: {exc}")
    return tables


def _detect_formulas_in_text(text: str) -> List[Dict[str, Any]]:
    """
    Обнаружить формулы в тексте по эвристическим шаблонам.

    Ищет:
      - LaTeX-формулы: $...$, $$...$$
      - Строки с характерными символами: интегралы, суммы, греческие буквы
    """
    formulas = []

    # LaTeX inline: $...$
    for match in re.finditer(r'\$([^$]+)\$', text):
        formulas.append({
            "content": match.group(1).strip(),
            "formula_type": "latex_inline",
            "start": match.start(),
            "end": match.end(),
        })

    # LaTeX block: $$...$$
    for match in re.finditer(r'\$\$(.+?)\$\$', text, re.DOTALL):
        formulas.append({
            "content": match.group(1).strip(),
            "formula_type": "latex_block",
            "start": match.start(),
            "end": match.end(),
        })

    # Подозреваемые формулы: строки с математическими символами
    math_pattern = re.compile(
        r'[=+\-*/^√∑∫∏∂∇∞≈≠≤≥±αβγδεζηθλμπρσφψω]'
    )
    for line in text.split('\n'):
        line = line.strip()
        if len(line) < 5:
            continue
        if math_pattern.search(line) and not line.startswith('|'):
            # Проверяем, что строка ещё не найдена как LaTeX
            already_found = any(
                f["start"] <= text.find(line) <= f["end"]
                for f in formulas
            )
            if not already_found:
                formulas.append({
                    "content": line,
                    "formula_type": "suspected_formula",
                    "start": text.find(line),
                    "end": text.find(line) + len(line),
                })

    return formulas


# ═══════════════════════════════════════════════════════════════
# Парсинг PDF
# ═══════════════════════════════════════════════════════════════
def parse_pdf(
    file_path: Path,
    model_hint: str = "",
) -> List[DocumentElement]:
    """
    Парсинг PDF-файла с извлечением текста, таблиц, формул и рисунков.

    Args:
        file_path:  Путь к PDF-файлу
        model_hint: Подсказка о модели (из имени файла)

    Returns:
        Список DocumentElement
    """
    if not file_path.exists():
        logger.error(f"Файл не найден: {file_path}")
        return []

    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []

    try:
        doc = fitz.open(str(file_path))
    except Exception as exc:
        logger.error(f"Не удалось открыть PDF {filename}: {exc}")
        return []

    logger.info(f"Парсинг PDF: {filename} ({len(doc)} стр.)")

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text").strip()

        # Проверка на скан
        if _is_scanned_page(page):
            pix = page.get_pixmap(dpi=OCR_DPI)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = _ocr_image(img)
            if ocr_text:
                elements.append(DocumentElement(
                    content=ocr_text,
                    element_type="text",
                    source=filename,
                    page=page_num + 1,
                    model=model,
                    metadata={"ocr": True},
                ))
            continue

        # Извлечение таблиц
        tables = _extract_tables_from_page(page)
        for table_data in tables:
            table_text = clean_text(table_data["text"])
            if table_text:
                detected_model = detect_model_in_text(table_text)
                elements.append(DocumentElement(
                    content=table_text,
                    element_type="table",
                    source=filename,
                    page=page_num + 1,
                    model=detected_model if detected_model != "unknown" else model,
                    metadata={
                        "num_rows": table_data["num_rows"],
                        "num_cols": table_data["num_cols"],
                        "bbox": table_data["bbox"],
                    },
                ))

        # Извлечение формул из текста страницы
        formulas = _detect_formulas_in_text(page_text)
        for formula_data in formulas:
            elements.append(DocumentElement(
                content=formula_data["content"],
                element_type="formula",
                source=filename,
                page=page_num + 1,
                model=model,
                metadata={"formula_type": formula_data["formula_type"]},
            ))

        # Извлечение изображений
        images = page.get_images(full=True)
        for img_idx, img_info in enumerate(images):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                img = Image.open(io.BytesIO(base_image["image"]))
                if img.width < MIN_IMAGE_WIDTH or img.height < MIN_IMAGE_HEIGHT:
                    continue

                # Поиск подписи к рисунку
                caption = find_figure_caption(page_text)
                content = caption if caption else f"Изображение {img_idx + 1}"

                # Сохранение изображения на диск
                FIGURES_DIR.mkdir(parents=True, exist_ok=True)
                img_filename = f"{file_path.stem}_p{page_num + 1}_img{img_idx}.png"
                img_path = FIGURES_DIR / img_filename
                img.save(str(img_path))

                elements.append(DocumentElement(
                    content=content,
                    element_type="figure",
                    source=filename,
                    page=page_num + 1,
                    model=model,
                    metadata={
                        "image_path": str(img_path),
                        "width": img.width,
                        "height": img.height,
                    },
                ))
            except Exception as exc:
                logger.debug(f"Ошибка извлечения изображения: {exc}")

        # Текст страницы (без таблиц и формул)
        if page_text:
            # Определяем заголовки
            headings = extract_headings(page_text)
            for heading in headings:
                detected_model = detect_model_in_text(heading)
                elements.append(DocumentElement(
                    content=heading,
                    element_type="heading",
                    source=filename,
                    page=page_num + 1,
                    model=detected_model if detected_model != "unknown" else model,
                ))

            # Основной текст
            cleaned = clean_text(page_text)
            if len(cleaned) >= OCR_MIN_TEXT_CHARS:
                detected_model = detect_model_in_text(cleaned)
                elements.append(DocumentElement(
                    content=cleaned,
                    element_type="text",
                    source=filename,
                    page=page_num + 1,
                    model=detected_model if detected_model != "unknown" else model,
                ))

    doc.close()
    logger.info(f"Извлечено {len(elements)} элементов из {filename}")
    return elements


# ═══════════════════════════════════════════════════════════════
# Парсинг DOCX
# ═══════════════════════════════════════════════════════════════
def parse_docx(
    file_path: Path,
    model_hint: str = "",
) -> List[DocumentElement]:
    """
    Парсинг DOCX-файла с извлечением текста и таблиц.

    Args:
        file_path:  Путь к DOCX-файлу
        model_hint: Подсказка о модели (из имени файла)

    Returns:
        Список DocumentElement
    """
    if not file_path.exists():
        logger.error(f"Файл не найден: {file_path}")
        return []

    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []

    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        logger.error(f"Не удалось открыть DOCX {filename}: {exc}")
        return []

    logger.info(f"Парсинг DOCX: {filename}")

    # Извлечение таблиц
    for table_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            continue

        # Формируем Markdown-таблицу
        num_cols = max(len(r) for r in rows_data)
        header = rows_data[0] if rows_data else []
        # Дополняем заголовки, если столбцов больше
        while len(header) < num_cols:
            header.append(f"Кол.{len(header) + 1}")

        md_lines = ["| " + " | ".join(header) + " |"]
        md_lines.append("| " + " | ".join(["---"] * num_cols) + " |")

        for row in rows_data[1:]:
            # Дополняем или обрезаем строку
            while len(row) < num_cols:
                row.append("")
            md_lines.append("| " + " | ".join(row[:num_cols]) + " |")

        table_md = "\n".join(md_lines)
        detected_model = detect_model_in_text(table_md)

        elements.append(DocumentElement(
            content=clean_text(table_md),
            element_type="table",
            source=filename,
            page=0,
            model=detected_model if detected_model != "unknown" else model,
            metadata={
                "num_rows": len(rows_data),
                "num_cols": num_cols,
                "table_index": table_idx,
            },
        ))

    # Извлечение текста из параграфов
    text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Проверяем, является ли параграф заголовком
        style_name = para.style.name.lower() if para.style else ""
        is_heading = "heading" in style_name or "заголовок" in style_name

        if is_heading:
            detected_model = detect_model_in_text(text)
            elements.append(DocumentElement(
                content=clean_text(text),
                element_type="heading",
                source=filename,
                page=0,
                model=detected_model if detected_model != "unknown" else model,
            ))
        else:
            text_parts.append(text)

    # Объединяем текст в один элемент
    if text_parts:
        combined = "\n\n".join(text_parts)
        cleaned = clean_text(combined)
        if len(cleaned) >= 50:
            detected_model = detect_model_in_text(cleaned)
            elements.append(DocumentElement(
                content=cleaned,
                element_type="text",
                source=filename,
                page=0,
                model=detected_model if detected_model != "unknown" else model,
            ))

    logger.info(f"Извлечено {len(elements)} элементов из {filename}")
    return elements
